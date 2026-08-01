"""DOCX extraction adapter, wrapping ``python-docx`` (extra ``viparse[office]``).

The engine walks the document body in order so headings, paragraphs, and tables
keep their original sequence, and records each run's **font name** as a signal —
that is how the S3 normalizer recognizes a legacy encoding (e.g. ``.VnTime`` →
TCVN3). It applies no Vietnamese logic itself.

``python-docx`` is imported lazily inside :meth:`DocxEngine.extract`, so importing
this module never requires the dependency; only extraction does.
"""

from __future__ import annotations

from collections.abc import Iterator
from typing import Any

from viparse.detect import CONTENT_TYPE_DOCX
from viparse.engines._shared import blocks_to_text
from viparse.errors import MissingDependency
from viparse.model import RawExtraction
from viparse.options import LoadOptions
from viparse.protocols import DEFAULT_PRIORITY, Source

_INSTALL_HINT = (
    "python-docx is required for DOCX extraction; install it with: pip install 'viparse[office]'"
)


def _import_docx() -> Any:
    """Import ``python-docx`` lazily, raising a clear error if it is missing."""
    try:
        import docx
    except ImportError as exc:
        raise MissingDependency(_INSTALL_HINT) from exc
    return docx


class DocxEngine:
    """Extracts ordered text, run fonts, and block structure from a ``.docx`` file."""

    priority = DEFAULT_PRIORITY
    #: Import name of the parse library this engine needs, and the extra that ships
    #: it — read by ``viparse doctor`` to report availability (``None`` = stdlib-only).
    dependency = "docx"
    extra = "office"

    def supports(self, content_type: str) -> bool:
        return content_type == CONTENT_TYPE_DOCX

    def extract(self, source: Source, options: LoadOptions) -> RawExtraction:
        docx = _import_docx()
        document = docx.Document(str(source))
        blocks: list[dict[str, Any]] = []
        fonts: set[str] = set()
        for kind, item in _iter_block_items(document):
            if kind == "paragraph":
                block_fonts: set[str] = set()
                _collect_fonts(item, block_fonts)
                fonts.update(block_fonts)
                block = _paragraph_block(item)
                if block is not None:
                    # Per-block font signal (SPEC-3 T3.2.4): lets the normalizer detect a
                    # mixed-encoding document and convert each block by its own encoding.
                    if block_fonts:
                        block["fonts"] = sorted(block_fonts)
                    blocks.append(block)
            else:  # table
                rows, table_fonts = _table_block(item)
                fonts.update(table_fonts)
                table_block: dict[str, Any] = {"type": "table", "rows": rows}
                if table_fonts:
                    table_block["fonts"] = sorted(table_fonts)
                blocks.append(table_block)
        signals: dict[str, Any] = {"fonts": sorted(fonts), "blocks": blocks}
        return RawExtraction(
            source=str(source),
            content_type=CONTENT_TYPE_DOCX,
            text=blocks_to_text(blocks),
            engine="docx",
            signals=signals,
        )


def _iter_block_items(document: Any) -> Iterator[tuple[str, Any]]:
    """Yield ``("paragraph"|"table", obj)`` for each body block, in document order.

    Descends into ``w:sdt`` content controls (Word form fields), whose blocks would
    otherwise be skipped — a common structure in Vietnamese government templates.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    p_tag, tbl_tag, sdt_tag, content_tag = (qn("w:p"), qn("w:tbl"), qn("w:sdt"), qn("w:sdtContent"))

    def walk(parent: Any) -> Iterator[tuple[str, Any]]:
        for child in parent.iterchildren():
            if child.tag == p_tag:
                yield "paragraph", Paragraph(child, document)
            elif child.tag == tbl_tag:
                yield "table", Table(child, document)
            elif child.tag == sdt_tag:
                content = child.find(content_tag)
                if content is not None:
                    yield from walk(content)

    yield from walk(document.element.body)


def _style_font(style: Any) -> str | None:
    """Return the font name declared on a paragraph/character style, if any."""
    if style is None:
        return None
    name: str | None = style.font.name
    return name


def _collect_fonts(paragraph: Any, fonts: set[str]) -> None:
    """Add every run's font name (the S3 encoding signal) to ``fonts``.

    Legacy documents often set the font once at the paragraph or character *style*
    level rather than per run, so ``run.font.name`` is ``None``; the style fonts are
    collected too, otherwise the encoding hint would be lost.
    """
    para_font = _style_font(paragraph.style)
    if para_font:
        fonts.add(para_font)
    for run in paragraph.runs:
        if run.font.name:
            fonts.add(run.font.name)
        run_style_font = _style_font(run.style)
        if run_style_font:
            fonts.add(run_style_font)


def _run_font(run: Any, paragraph: Any) -> str | None:
    """The effective font for a run: its own, else its character style, else the paragraph's.

    Mirrors the inheritance in :func:`_collect_fonts` but resolved *per run*, so the S3
    normalizer can detect a paragraph that mixes a legacy-font run with a Unicode one.
    """
    if run.font.name:
        name: str = run.font.name
        return name
    return _style_font(run.style) or _style_font(paragraph.style)


# OOXML encodes a soft hyphen as an *element*, ``<w:softHyphen/>``, not as U+00AD in a
# ``<w:t>``. python-docx's ``run.text`` concatenates only text nodes, tabs and breaks, so
# the character disappears.
#
# For most documents that loses a typographic hint. For TCVN3 it loses a letter: 0xAD is
# ``ư``, so a legacy .doc routed through LibreOffice's docx conversion came back with every
# ``ư`` gone — được, người, trường, nước all silently missing a character. Measured on ten
# real government documents: 848 of 848 occurrences lost.
_SOFT_HYPHEN = "\u00ad"


def _run_text(run: Any) -> str:
    """``run.text``, but keeping the soft hyphen.

    Mirrors python-docx's own mapping (``w:t`` → text, ``w:tab`` → tab, ``w:br``/``w:cr``
    → newline) and adds ``w:softHyphen`` → U+00AD.
    """
    from docx.oxml.ns import qn

    t_tag, tab_tag, br_tag, cr_tag, shy_tag = (
        qn("w:t"),
        qn("w:tab"),
        qn("w:br"),
        qn("w:cr"),
        qn("w:softHyphen"),
    )
    parts: list[str] = []
    for node in run._element.iter():
        if node.tag == t_tag:
            parts.append(node.text or "")
        elif node.tag == tab_tag:
            parts.append("\t")
        elif node.tag in (br_tag, cr_tag):
            parts.append("\n")
        elif node.tag == shy_tag:
            parts.append(_SOFT_HYPHEN)
    return "".join(parts)


# Elements that wrap runs inside a paragraph. `paragraph.runs` returns only *direct*
# `w:r` children, so a run inside any of these is invisible to it — in one real
# government document 10,473 of 25,377 characters, 41%, sat inside `w:ins`.
_RUN_CONTAINERS = ("w:ins", "w:hyperlink", "w:smartTag", "w:moveTo", "w:sdtContent")

# ...and the ones whose runs are *not* part of the document text. `w:del` and
# `w:moveFrom` hold what a tracked change removed; including them would resurrect
# deleted sentences, which is a worse failure than dropping inserted ones.
_REMOVED_CONTAINERS = ("w:del", "w:moveFrom")


def _paragraph_run_elements(paragraph: Any) -> list[Any]:
    """Every ``w:r`` in the paragraph, in document order, minus deleted ones.

    ``paragraph.runs`` is only the direct children. A document that has been through
    review carries most of its prose inside ``w:ins`` — LibreOffice writes tracked
    insertions when converting a legacy ``.doc`` that has them — and none of it reaches
    the text.

    This is the same shape as the soft-hyphen loss (VIP-87): a convenience accessor that
    under-reports, silently, on real documents but never on a synthetic fixture.
    """
    from docx.oxml.ns import qn

    removed = {qn(tag) for tag in _REMOVED_CONTAINERS}
    root = paragraph._p
    kept: list[Any] = []
    for run in root.iter(qn("w:r")):
        node = run.getparent()
        while node is not None and node is not root:
            if node.tag in removed:
                break
            node = node.getparent()
        else:
            kept.append(run)
    return kept


def _paragraph_run_objects(paragraph: Any) -> list[Any]:
    """:func:`_paragraph_run_elements` wrapped as python-docx ``Run`` objects.

    ``_run_font`` reads run properties and style inheritance through the wrapper, so the
    elements alone are not enough.
    """
    from docx.text.run import Run

    return [Run(element, paragraph) for element in _paragraph_run_elements(paragraph)]


def _paragraph_text(paragraph: Any) -> str:
    """``paragraph.text`` built from :func:`_run_text`.

    python-docx defines paragraph text as the concatenation of its run texts, so building
    it the same way keeps the block text and the run segments in step. Fixing only one of
    them would shift the run boundaries the normalizer relies on for per-run font
    detection — which is why both go through :func:`_paragraph_run_objects` rather than
    ``paragraph.runs``.
    """
    return "".join(_run_text(run) for run in _paragraph_run_objects(paragraph))


def _cell_text(cell: Any) -> str:
    """``cell.text`` built from :func:`_paragraph_text`.

    Table cells are a second path to the same loss: python-docx joins paragraph texts
    here too, so fixing only the paragraph path left 23 of 848 soft hyphens missing —
    the ones inside tables, which is where Vietnamese administrative forms put most of
    their content.
    """
    return "\n".join(_paragraph_text(paragraph) for paragraph in cell.paragraphs)


def _paragraph_runs(paragraph: Any) -> list[dict[str, Any]]:
    """Per-run ``{text, font}`` segments whose texts concatenate back to ``paragraph.text``.

    Empty runs are dropped (they contribute nothing to the text); this keeps the segment
    list faithful so the normalizer can trust it for per-run conversion.
    """
    return [
        {"text": text, "font": _run_font(run, paragraph)}
        for run, text in ((run, _run_text(run)) for run in _paragraph_run_objects(paragraph))
        if text
    ]


def _paragraph_block(paragraph: Any) -> dict[str, Any] | None:
    """Map a paragraph to a heading/paragraph block, or ``None`` if empty."""
    text = _paragraph_text(paragraph)
    style = paragraph.style.name if paragraph.style is not None else ""
    if style.startswith("Heading"):
        block: dict[str, Any] = {"type": "heading", "level": _heading_level(style), "text": text}
    elif not text.strip():
        return None
    else:
        block = {"type": "paragraph", "text": text}
    # Per-run font signal (SPEC-3 T3.2.4): lets the normalizer convert a paragraph that
    # mixes a legacy-font run with a Unicode one at run granularity, not whole-block.
    runs = _paragraph_runs(paragraph)
    if runs:
        block["runs"] = runs
    return block


def _heading_level(style_name: str) -> int:
    """Parse the level from a heading style name (``"Heading 2"`` → 2)."""
    tail = style_name.split()[-1]
    return int(tail) if tail.isdigit() else 1


def _table_block(table: Any) -> tuple[list[list[str]], set[str]]:
    """Extract a table's cell text (rows × cols) and any run fonts within it.

    Horizontally-merged cells share one underlying element and would otherwise be
    read once per spanned column; consecutive repeats are collapsed so the value
    appears a single time.
    """
    rows: list[list[str]] = []
    fonts: set[str] = set()
    for row in table.rows:
        cells: list[str] = []
        previous_tc: Any = None
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _collect_fonts(paragraph, fonts)
            if cell._tc is previous_tc:
                continue  # part of a horizontal merge already captured
            previous_tc = cell._tc
            cells.append(_cell_text(cell))
        rows.append(cells)
    return rows, fonts
