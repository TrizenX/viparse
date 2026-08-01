"""Tests for the DOCX extraction adapter."""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

from viparse.detect import CONTENT_TYPE_DOCX
from viparse.engines.docx import DocxEngine
from viparse.errors import MissingDependency
from viparse.options import LoadOptions
from viparse.registry import EngineRegistry

docx = pytest.importorskip("docx")  # python-docx; skipped if the office extra is absent


def _make_docx(path: Path) -> Path:
    document = docx.Document()
    document.add_heading("Tiêu đề", level=1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Tiếng Việt")
    run.font.name = ".VnTime"
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    document.add_paragraph("cuối")
    document.save(str(path))
    return path


def test_extract_returns_raw_extraction(tmp_path: Path) -> None:
    raw = DocxEngine().extract(_make_docx(tmp_path / "a.docx"), LoadOptions())
    assert raw.engine == "docx"
    assert raw.content_type == CONTENT_TYPE_DOCX
    assert "Tiếng Việt" in raw.text
    assert "cuối" in raw.text


def test_extract_captures_run_font_signal(tmp_path: Path) -> None:
    """The MVP acceptance: a run's font name is surfaced for the S3 detector."""
    raw = DocxEngine().extract(_make_docx(tmp_path / "a.docx"), LoadOptions())
    assert ".VnTime" in raw.signals["fonts"]


def test_extract_preserves_block_order(tmp_path: Path) -> None:
    raw = DocxEngine().extract(_make_docx(tmp_path / "a.docx"), LoadOptions())
    kinds = [block["type"] for block in raw.signals["blocks"]]
    assert kinds == ["heading", "paragraph", "table", "paragraph"]
    heading = raw.signals["blocks"][0]
    assert heading["level"] == 1
    assert heading["text"] == "Tiêu đề"


def test_extract_table_structure(tmp_path: Path) -> None:
    raw = DocxEngine().extract(_make_docx(tmp_path / "a.docx"), LoadOptions())
    table = next(block for block in raw.signals["blocks"] if block["type"] == "table")
    assert table["rows"] == [["A", "B"], ["C", "D"]]
    assert "A\tB" in raw.text


def test_empty_paragraphs_are_skipped(tmp_path: Path) -> None:
    document = docx.Document()
    document.add_paragraph("")  # blank spacer
    document.add_paragraph("content")
    path = tmp_path / "b.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    assert [block["text"] for block in raw.signals["blocks"]] == ["content"]


def test_captures_style_inherited_font(tmp_path: Path) -> None:
    """A legacy font set via a character style (not per-run) is still surfaced."""
    from docx.enum.style import WD_STYLE_TYPE

    document = docx.Document()
    style = document.styles.add_style("Legacy", WD_STYLE_TYPE.CHARACTER)
    style.font.name = "VNI-Times"
    run = document.add_paragraph().add_run("styled text")
    run.style = document.styles["Legacy"]
    assert run.font.name is None  # font is only on the style, not the run
    path = tmp_path / "styled.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    assert "VNI-Times" in raw.signals["fonts"]


def test_captures_paragraph_style_font(tmp_path: Path) -> None:
    """A font set on the paragraph style (not per-run) is surfaced too."""
    from docx.enum.style import WD_STYLE_TYPE

    document = docx.Document()
    style = document.styles.add_style("LegacyPara", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "VNI-Para"
    document.add_paragraph("body", style="LegacyPara")
    path = tmp_path / "parastyle.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    assert "VNI-Para" in raw.signals["fonts"]


def test_style_font_helper_handles_missing_style() -> None:
    from viparse.engines.docx import _style_font

    assert _style_font(None) is None


def test_empty_content_control_is_ignored(tmp_path: Path) -> None:
    """A malformed w:sdt with no sdtContent is skipped without crashing."""
    from docx.oxml import OxmlElement

    document = docx.Document()
    document.add_paragraph("real content")
    document.element.body.append(OxmlElement("w:sdt"))  # no sdtContent child
    path = tmp_path / "emptysdt.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    assert "real content" in raw.text


def test_captures_table_cell_font(tmp_path: Path) -> None:
    document = docx.Document()
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    cell.paragraphs[0].add_run("cell").font.name = ".VnArial"
    path = tmp_path / "cellfont.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    assert ".VnArial" in raw.signals["fonts"]


def test_horizontally_merged_cells_are_not_duplicated(tmp_path: Path) -> None:
    document = docx.Document()
    table = document.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "merged"
    table.cell(0, 2).text = "third"
    table.cell(0, 0).merge(table.cell(0, 1))
    path = tmp_path / "merged.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    table_block = next(b for b in raw.signals["blocks"] if b["type"] == "table")
    assert table_block["rows"] == [["merged", "third"]]


def test_content_control_paragraphs_are_not_dropped(tmp_path: Path) -> None:
    """Text inside a w:sdt content control is extracted, not silently lost."""
    from docx.oxml import OxmlElement

    document = docx.Document()
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    para = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "inside control"
    run.append(text)
    para.append(run)
    content.append(para)
    sdt.append(content)
    document.element.body.append(sdt)
    path = tmp_path / "sdt.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    assert "inside control" in raw.text


def test_attaches_per_block_font_signal(tmp_path: Path) -> None:
    """Each block carries its own font signal so the S3 normalizer can detect mixing."""
    raw = DocxEngine().extract(_make_docx(tmp_path / "a.docx"), LoadOptions())
    para = next(b for b in raw.signals["blocks"] if b.get("text") == "Tiếng Việt")
    assert ".VnTime" in para["fonts"]


def test_mixed_encoding_docx_converts_per_block_end_to_end(tmp_path: Path) -> None:
    """A .VnTime (TCVN3) paragraph and a Unicode paragraph in one file (SPEC-3 T3.2.4).

    The Unicode paragraph legitimately contains "®"; whole-document conversion would
    turn it into "đ". Per-block detection must convert only the legacy paragraph.
    """
    from viparse.normalize.normalizer import VietnameseNormalizer

    document = docx.Document()
    legacy = document.add_paragraph().add_run("®¸")  # TCVN3 surface for "đá"
    legacy.font.name = ".VnTime"
    unicode_run = document.add_paragraph().add_run("viparse® 2026")  # already Unicode
    unicode_run.font.name = "Arial"
    path = tmp_path / "mixed.docx"
    document.save(str(path))

    raw = DocxEngine().extract(path, LoadOptions())
    nd = VietnameseNormalizer().normalize(raw, LoadOptions())
    assert "đá" in nd.text  # the legacy paragraph converted
    assert "viparse® 2026" in nd.text  # the Unicode "®" is preserved, not mapped to "đ"
    assert nd.encoding_detected == "tcvn3"


def test_attaches_per_run_font_signal(tmp_path: Path) -> None:
    """A paragraph's runs each carry their own font, in order (SPEC-3 T3.2.4, per-run)."""
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("abc").font.name = ".VnTime"
    paragraph.add_run("def").font.name = "Arial"
    path = tmp_path / "runs.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    block = next(b for b in raw.signals["blocks"] if b.get("type") == "paragraph")
    assert block["runs"] == [
        {"text": "abc", "font": ".VnTime"},
        {"text": "def", "font": "Arial"},
    ]


def test_mixed_run_paragraph_converts_per_run_end_to_end(tmp_path: Path) -> None:
    """A single paragraph mixing a .VnTime run and a Unicode run (SPEC-3 T3.2.4, VIP-72).

    Whole-block conversion would corrupt the Unicode run's "®"; per-run detection must
    convert only the legacy run.
    """
    from viparse.normalize.normalizer import VietnameseNormalizer

    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("®¸ ").font.name = ".VnTime"  # TCVN3 surface for "đá"
    paragraph.add_run("viparse® 2026").font.name = "Arial"  # already Unicode
    path = tmp_path / "mixedrun.docx"
    document.save(str(path))

    raw = DocxEngine().extract(path, LoadOptions())
    nd = VietnameseNormalizer().normalize(raw, LoadOptions())
    assert "đá" in nd.text  # the legacy run converted
    assert "viparse® 2026" in nd.text  # the Unicode "®" is preserved, not mapped to "đ"
    assert nd.encoding_detected == "tcvn3"


def test_empty_heading_carries_no_run_signal(tmp_path: Path) -> None:
    """An empty heading yields a heading block with no per-run signal (not skipped)."""
    document = docx.Document()
    document.add_heading("", level=1)
    document.add_paragraph("body")
    path = tmp_path / "emptyheading.docx"
    document.save(str(path))
    raw = DocxEngine().extract(path, LoadOptions())
    heading = next(b for b in raw.signals["blocks"] if b.get("type") == "heading")
    assert "runs" not in heading


def test_supports_only_docx() -> None:
    engine = DocxEngine()
    assert engine.supports(CONTENT_TYPE_DOCX)
    assert not engine.supports("application/pdf")


def test_registry_selects_docx_engine() -> None:
    reg = EngineRegistry()
    reg.register(DocxEngine())
    assert isinstance(reg.select(CONTENT_TYPE_DOCX), DocxEngine)


def test_missing_dependency_raises_clear_error(monkeypatch: pytest.MonkeyPatch) -> None:
    """Without python-docx, extraction fails with actionable install guidance."""
    monkeypatch.setitem(sys.modules, "docx", None)
    with pytest.raises(MissingDependency, match=r"viparse\[office\]"):
        DocxEngine().extract("missing.docx", LoadOptions())


# --- Soft hyphen: a letter in TCVN3, not decoration (VIP-87) ---------------------------


def _append_soft_hyphen(run: object) -> None:
    """Append ``<w:softHyphen/>`` to a run.

    Written through the XML rather than by typing U+00AD into the text, because Word and
    LibreOffice both store the character as an element and never as a text node. A
    fixture built the convenient way would contain a literal U+00AD, survive extraction
    with or without the fix, and prove nothing.
    """
    from docx.oxml.ns import qn

    element = run._element  # type: ignore[attr-defined]
    element.append(element.makeelement(qn("w:softHyphen"), {}))


def test_soft_hyphen_survives_paragraph_extraction(tmp_path: Path) -> None:
    """0xAD is ``ư`` in TCVN3, so dropping it deletes a letter, not a hint.

    python-docx's ``run.text`` ignores ``<w:softHyphen/>``. Routed through LibreOffice's
    .doc → .docx conversion, that lost every ``ư`` in a legacy document — 848 of 848
    across a ten-document corpus, silently.
    """
    docx = pytest.importorskip("docx")
    document = docx.Document()
    run = document.add_paragraph().add_run("ng")
    _append_soft_hyphen(run)
    source = tmp_path / "shy_paragraph.docx"
    document.save(str(source))

    extraction = DocxEngine().extract(source, LoadOptions())
    assert "\u00ad" in extraction.text


def test_soft_hyphen_survives_table_extraction(tmp_path: Path) -> None:
    """Table cells are a second path to the same loss.

    Fixing only the paragraph path left the soft hyphens inside tables missing, which is
    where Vietnamese administrative forms keep most of their content.
    """
    docx = pytest.importorskip("docx")
    document = docx.Document()
    table = document.add_table(rows=1, cols=1)
    run = table.cell(0, 0).paragraphs[0].add_run("t")
    _append_soft_hyphen(run)
    source = tmp_path / "shy_table.docx"
    document.save(str(source))

    extraction = DocxEngine().extract(source, LoadOptions())
    assert "\u00ad" in extraction.text


# --- Tracked insertions: body text python-docx cannot see (VIP-95) ---------------------


def _wrap_run_in(paragraph: object, text: str, tag: str) -> None:
    """Move a new run into a ``<{tag}>`` wrapper inside the paragraph.

    Built through the XML because that is how the loss happens: `paragraph.runs` returns
    only *direct* `w:r` children, so a fixture that adds the run normally would be
    reachable with or without the fix and would prove nothing.
    """
    from docx.oxml.ns import qn

    run = paragraph.add_run(text)  # type: ignore[attr-defined]
    element = run._element
    wrapper = element.makeelement(qn(tag), {})
    element.getparent().replace(element, wrapper)
    wrapper.append(element)


def test_tracked_insertion_text_is_extracted(tmp_path: Path) -> None:
    """``w:ins`` holds ordinary prose, and it was being dropped entirely.

    In one real government document 10,473 of 25,377 characters — 41% — sat inside
    ``w:ins``, because LibreOffice writes tracked insertions when converting a legacy
    ``.doc`` that carries them. None of it reached the text.
    """
    docx = pytest.importorskip("docx")
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("visible ")
    _wrap_run_in(paragraph, "inserted", "w:ins")
    source = tmp_path / "tracked_insert.docx"
    document.save(str(source))

    extraction = DocxEngine().extract(source, LoadOptions())
    assert "inserted" in extraction.text
    assert "visible inserted" in extraction.text


def test_tracked_deletion_text_is_not_extracted(tmp_path: Path) -> None:
    """``w:del`` is what a tracked change *removed*.

    Reaching into wrappers indiscriminately would resurrect deleted sentences, which is
    a worse failure than dropping inserted ones — the text would be wrong rather than
    short, and nothing downstream could tell.
    """
    docx = pytest.importorskip("docx")
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("kept ")
    _wrap_run_in(paragraph, "removed", "w:del")
    source = tmp_path / "tracked_delete.docx"
    document.save(str(source))

    assert "removed" not in DocxEngine().extract(source, LoadOptions()).text


def test_tracked_insertion_appears_in_run_segments(tmp_path: Path) -> None:
    """The per-run font signal must see the same runs the text does.

    The normalizer converts mixed-encoding paragraphs at run granularity, so a run that
    is in the text but not in the segment list would be converted with the wrong table.
    """
    docx = pytest.importorskip("docx")
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("a")
    _wrap_run_in(paragraph, "b", "w:ins")
    source = tmp_path / "tracked_runs.docx"
    document.save(str(source))

    blocks = DocxEngine().extract(source, LoadOptions()).signals["blocks"]
    segments = [seg["text"] for block in blocks for seg in block.get("runs", [])]
    assert segments == ["a", "b"]


def test_hyperlink_wrapped_run_is_extracted(tmp_path: Path) -> None:
    """``w:hyperlink`` wraps runs the same way, and is far more common than tracked changes."""
    docx = pytest.importorskip("docx")
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("see ")
    _wrap_run_in(paragraph, "the site", "w:hyperlink")
    source = tmp_path / "hyperlink.docx"
    document.save(str(source))

    assert "see the site" in DocxEngine().extract(source, LoadOptions()).text
